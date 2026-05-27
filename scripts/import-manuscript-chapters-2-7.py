from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
CONTENT = ROOT / "content"
KJV_PATH = Path("/private/tmp/revelation-kjv.json")

DOCS = {
    2: Path("/Users/samuel/Downloads/Revelation Chapter Two.docx"),
    3: Path("/Users/samuel/Downloads/Revelation Chapter Three.docx"),
    4: Path("/Users/samuel/Downloads/Revelation Chapter Four.docx"),
    5: Path("/Users/samuel/Downloads/Revelation Chapter Five.docx"),
    6: Path("/Users/samuel/Downloads/Revelation Chapter Six.docx"),
    7: Path("/Users/samuel/Downloads/Revelation Chapter Seven.docx"),
}

DOC_SOURCE_IDS = {
    2: "revelation-chapter-two-docx",
    3: "revelation-chapter-three-docx",
    4: "revelation-chapter-four-docx",
    5: "revelation-chapter-five-docx",
    6: "revelation-chapter-six-docx",
    7: "revelation-chapter-seven-docx",
}

MCNULTY_ID = "revelation-practical-living-in-the-judgment-hour"
SHORTER_ID = "revelation-a-shorter-commentary"


def words(text: str) -> int:
    return len([part for part in text.strip().split() if part])


def clean_doc_text(text: str) -> str:
    text = re.sub(r"filecite[^]*", "", text)
    text = re.sub(r"\([^)]*(?:PDF|pp\.|p\.|uploaded|McNulty|Beale|Osborne|Stefanovi|Maxwell)[^)]*\)", "", text)
    text = re.sub(r"\b(?:McNulty|Maxwell|Stefanović|Stefanovic|Beale|Osborne|Akin|deSilva|Poythress)\s+(?:argues|says|notes|emphasizes|reads|treats|connects)\b", "the interpretation explains", text)
    return re.sub(r"\s+", " ", text).strip()


def read_doc_headings(path: Path) -> list[str]:
    document = Document(str(path))
    headings = []
    for paragraph in document.paragraphs:
        text = clean_doc_text(paragraph.text.strip())
        if text and paragraph.style.name.startswith("Heading"):
            headings.append(text)
    return headings


SOURCE_NAMES = (
    "McNulty",
    "Maxwell",
    "Stefanovic",
    "Stefanović",
    "Osborne",
    "Beale",
    "Campbell",
    "Akin",
    "deSilva",
    "Poythress",
    "Amazing Facts",
    "Revelation Verse by Verse",
    "Ellen White",
)

DOC_UNIT_CACHE: dict[int, dict[int, str]] = {}


def parse_ref_range(text: str, chapter: int) -> tuple[int, int] | None:
    pattern = rf"^(?:Verse / phrase\s+—\s+)?(?:Rev\.?|Revelation)\s+{chapter}:(\d+)(?:\s*[–-]\s*(\d+))?"
    match = re.match(pattern, text)
    if not match:
        return None
    start = int(match.group(1))
    end = int(match.group(2) or start)
    return start, end


def add_doc_unit(units: dict[int, list[str]], verses: tuple[int, int] | None, parts: list[str]) -> None:
    if not verses or not parts:
        return
    text = " ".join(parts).strip()
    if not text:
        return
    start, end = verses
    for verse in range(start, end + 1):
        units.setdefault(verse, []).append(text)


def read_doc_units(chapter: int) -> dict[int, str]:
    if chapter in DOC_UNIT_CACHE:
        return DOC_UNIT_CACHE[chapter]
    document = Document(str(DOCS[chapter]))
    units: dict[int, list[str]] = {}
    current_verses: tuple[int, int] | None = None
    current_parts: list[str] = []
    skip_section = False

    for paragraph in document.paragraphs:
        text = clean_doc_text(paragraph.text.strip())
        if not text:
            continue
        is_heading = paragraph.style.name.startswith("Heading")
        if is_heading:
            add_doc_unit(units, current_verses, current_parts)
            current_verses = None
            current_parts = []
            lowered = text.lower()
            skip_section = (
                "bibliography" in lowered
                or "open questions" in lowered
                or "limitations" in lowered
                or text.startswith("Revelation Chapter")
            )
            continue
        if skip_section:
            continue
        found_range = parse_ref_range(text, chapter)
        if found_range:
            add_doc_unit(units, current_verses, current_parts)
            current_verses = found_range
            current_parts = [text]
            continue
        if current_verses:
            current_parts.append(text)

    add_doc_unit(units, current_verses, current_parts)
    DOC_UNIT_CACHE[chapter] = {verse: " ".join(parts) for verse, parts in units.items()}
    return DOC_UNIT_CACHE[chapter]


def public_doc_text(text: str) -> str:
    text = clean_doc_text(text)
    text = re.sub(r"^(?:Verse / phrase\s+—\s+)?(?:Rev\.?|Revelation)\s+\d+:\d+(?:\s*[–-]\s*\d+)?\.?,?\s*", "", text)
    text = re.sub(r"\b(?:Observation|Interpretation|Application|Symbols and imagery|Theological significance|Cross-references|Passage summary|Key takeaway)\.\s*", "", text)
    text = re.sub(r"\b(?:Selected bibliography|interpretive frame|Page references below).*?(?=Revelation|Verse / phrase|$)", "", text, flags=re.I)
    text = re.sub(r"\b(?:uploaded|source|sources|resource|resources|PDF|pagination|page references?)\b[^.?!]*(?:[.?!]|$)", "", text, flags=re.I)
    adventist_names = r"(?:McNulty|Maxwell|Revelation Verse by Verse|Amazing Facts)"
    academic_names = r"(?:Stefanovic|Stefanović|Osborne|Beale(?: and Campbell)?|Campbell|Akin|deSilva|Poythress)"
    verbs = r"(reads?|identif(?:y|ies)|argues?|stress(?:es)?|emphasiz(?:e|es)|notes?|connects?|appl(?:y|ies)|takes?|sees?|treats?|describes?|highlights?|explains?|echo(?:es)?|ties?|regards?|reports?)"
    text = re.sub(rf"\b{adventist_names}(?:\s*(?:,|and)\s*{adventist_names})*\s+{verbs}\b", r"Adventist interpretation \1", text)
    text = re.sub(rf"\b{academic_names}(?:\s*(?:,|and)\s*{academic_names})*\s+{verbs}\b", r"commentators \1", text)
    text = text.replace("the study material explains", "the interpretation explains")
    text = text.replace("the interpretation explains it as", "it presents it as")
    text = text.replace("the interpretation explains that", "the interpretation emphasizes that")
    text = text.replace("and echo the same", "Adventist explanations echo the same")
    text = text.replace("and commentators tie", "Other commentators tie")
    text = text.replace("and likewise argue", "Other commentators likewise argue")
    text = text.replace("For, then,", "For this reading, then,")
    text = text.replace("makes the same point", "Adventist interpretation makes the same point")
    text = text.replace("ends his chapter summary", "A concise theological summary ends")
    for name in SOURCE_NAMES:
        text = re.sub(rf"\b{re.escape(name)}(?:'s|’s)?\b", "", text)
    text = text.replace("Adventist interpretation echo ", "Adventist interpretation echoes ")
    text = text.replace("Other commentators both highlight", "Other commentators highlight")
    text = text.replace("when he says that", "by stressing that")
    text = text.replace("and he then", "and this reading then")
    text = text.replace("he then", "this reading then")
    text = text.replace("He ties this", "This connects the point")
    text = text.replace("He also stresses", "The point also stresses")
    text = text.replace("his most important theological moves", "the most important theological points")
    text = text.replace("For, then,", "For this reading, then,")
    text = text.replace("Adds a crucial", "This adds a crucial")
    text = text.replace("The interpretation explains", "This reading explains")
    text = text.replace("The interpretation emphasizes", "This reading emphasizes")
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def split_sentences(text: str) -> list[str]:
    sentences = [part.strip() for part in re.split(r"(?<=[.!?])\s+", text) if len(part.strip().split()) >= 5]
    if sentences and sentences[0].startswith(("“", '"')) and "…" in sentences[0]:
        repaired = re.sub(r"^[“\"][^”\"]+[”\"]\s*", "", sentences[0]).strip()
        if len(repaired.split()) >= 5:
            sentences[0] = repaired
        else:
            sentences = sentences[1:]
    if sentences and sentences[0][0].islower() and "…" in sentences[0]:
        sentences = sentences[1:]
    elif sentences and sentences[0][0].islower() and len(sentences[0].split()) < 18:
        sentences = sentences[1:]
    normalized = []
    for sentence in sentences:
        sentence = sentence.strip()
        lowered = sentence.lower()
        if any(
            fragment in lowered
            for fragment in [
                "pdf",
                "uploaded",
                "cited",
                "source",
                "found in and",
                ",,",
                "drawing on as",
                "consulted",
            ]
        ):
            continue
        if re.match(r"^(?:rev\.?\s*)?\d+:\d+", lowered):
            continue
        if lowered.startswith(("he ", "his ", "gives ", "explicitly ", "disagrees ", "similarly ", "binds ", "makes ", "also ")):
            continue
        if "his reading" in lowered or "his discussion" in lowered:
            continue
        if sentence.startswith("and "):
            sentence = "Other commentators " + sentence[4:]
        if sentence and sentence[0].islower():
            sentence = sentence[0].upper() + sentence[1:]
        sentence = sentence.replace("The interpretation emphasizes", "This reading emphasizes")
        sentence = sentence.replace("The interpretation explains", "This reading explains")
        sentence = sentence.replace("Other commentators both highlight", "Other commentators highlight")
        sentence = sentence.replace("For, then,", "For this reading, then,")
        sentence = sentence.replace("this reading then explicitly", "this reading explicitly")
        sentence = sentence.replace("It presents it as", "It presents the promise as")
        if sentence.startswith("Adds a crucial"):
            sentence = "This adds a crucial" + sentence[len("Adds a crucial"):]
        normalized.append(sentence)
    return normalized


def public_guard(text: str) -> str:
    text = text.replace("belongs to", "is held by")
    text = text.replace("Belongs to", "Is held by")
    text = text.replace("salvation is held by God and to the Lamb", "salvation comes from God and from the Lamb")
    text = text.replace("the throne is held by God alone", "the throne is God's alone")
    text = text.replace("the right to ultimate worship is held by Him because creation itself is held by Him", "the right to ultimate worship is due to Him because He created all things")
    text = text.replace("The interpretation", "This reading")
    text = text.replace("This reading explains the seal", "The seal is explained")
    text = text.replace("This reading emphasizes that the seal", "The seal")
    text = text.replace("and this reading explicitly ties", "and is explicitly tied")
    text = text.replace("Adventist interpretation echoes the same broad understanding", "Other Adventist explanations echo the same broad understanding")
    text = text.replace("Adds an important", "An important")
    text = text.replace("Also reports the historicist identification", "The historicist identification")
    text = text.replace("The historicist identification of these signs and situates them", "This historicist identification situates these signs")
    text = text.replace("while the symbolic language also reminds us that the language John uses is never", "while the symbolic language also reminds us that John's imagery is never")
    text = text.replace("doctrinal correctives", "doctrinal correction")
    text = text.replace("while rightly reminds us", "while the symbolic language also reminds us")
    text = text.replace("non-Adventist treatment helps illuminate", "historical background helps illuminate")
    text = text.replace("That reading is intentionally sharp should not be softened away", "The diagnosis is intentionally sharp and should not be softened")
    text = text.replace("That reading is intentionally sharp and should not be softened away", "The diagnosis is intentionally sharp and should not be softened")
    text = text.replace("Key takeaway:", "")
    text = re.sub(r"\s+", " ", text).strip() if "\n\n" not in text else "\n\n".join(re.sub(r"\s+", " ", p).strip() for p in text.split("\n\n"))
    return text


def select_doc_excerpt(chapter: int, verse: int, bible_text: str) -> str:
    raw = read_doc_units(chapter).get(verse, "")
    if not raw:
        return ""
    text = public_doc_text(raw)
    sentences = split_sentences(text)
    if not sentences:
        return ""
    terms = [term.lower() for term in key_terms(bible_text)]
    terms.extend(part.lower().strip(" ,.;:!?\"'()") for part in bible_text.split() if len(part.strip(" ,.;:!?\"'()")) > 5)
    priority: list[str] = []
    secondary: list[str] = []
    for sentence in sentences:
        lowered = sentence.lower()
        if any(term and term in lowered for term in terms):
            priority.append(sentence)
        else:
            secondary.append(sentence)
    ordered = priority + secondary
    chosen: list[str] = []
    for sentence in ordered:
        trial = " ".join(chosen + [sentence])
        if words(trial) <= 235:
            chosen.append(sentence)
        if words(" ".join(chosen)) >= 150:
            break
    if words(" ".join(chosen)) < 120:
        for sentence in secondary:
            if sentence in chosen:
                continue
            trial = " ".join(chosen + [sentence])
            if words(trial) <= 255:
                chosen.append(sentence)
            if words(" ".join(chosen)) >= 145:
                break
    return public_guard(" ".join(chosen).strip())


def ref(source_id: str, locator: str, claim_type: str, priority: int) -> dict:
    return {
        "sourceId": source_id,
        "locator": locator,
        "claimType": claim_type,
        "priority": priority,
    }


CHAPTER_META = {
    2: {
        "title": "Messages to Ephesus, Smyrna, Pergamos, and Thyatira",
        "summary": "Revelation 2 begins the messages to the seven churches. Christ walks among His lampstands and speaks to churches marked by love grown cold, suffering faithfulness, doctrinal compromise, and tolerated corruption.",
        "historical": "The four churches were real congregations in Roman Asia, yet the messages also speak to the church across history. In Adventist historicist interpretation, Ephesus reflects the apostolic church, Smyrna the persecuted church, Pergamos the age of compromise after imperial favor, and Thyatira the long medieval period with both corruption and faithful witness.",
        "literary": "The chapter continues the vision of Revelation 1. Each message begins with a description of Christ, gives diagnosis, offers counsel or warning, and closes with a promise to the overcomer.",
        "themes": ["Christ among the churches", "First love", "Persecution", "Compromise", "Nicolaitanes", "Balaam", "Jezebel", "Overcoming"],
        "outline": [
            ("2:1-7", "Ephesus", "Truth without first love must remember, repent, and recover living devotion."),
            ("2:8-11", "Smyrna", "A suffering church receives no rebuke and is called to faithfulness unto death."),
            ("2:12-17", "Pergamos", "A loyal but compromised church must reject Balaam-like and Nicolaitan compromise."),
            ("2:18-29", "Thyatira", "A growing church must stop tolerating Jezebel-like teaching and hold fast until Christ comes."),
        ],
        "periods": {
            "Ephesus": "Adventist historicist application commonly sees Ephesus as the apostolic and immediately post-apostolic church: doctrinally alert, missionary, and yet vulnerable to cooling love.",
            "Smyrna": "Smyrna is commonly connected with the persecuted church under pagan Rome, where poverty and martyrdom could not erase spiritual riches.",
            "Pergamos": "Pergamos is commonly connected with the period of imperial favor and compromise, when worldly advantage threatened the purity of Christian witness.",
            "Thyatira": "Thyatira is commonly connected with the medieval period, where corrupt religious influence grew while faithful believers were still known by Christ.",
        },
    },
    3: {
        "title": "Messages to Sardis, Philadelphia, and Laodicea",
        "summary": "Revelation 3 completes the seven messages. Sardis has a reputation for life while dying, Philadelphia has little strength yet an open door, and Laodicea is prosperous, lukewarm, and urgently invited back into fellowship with Christ.",
        "historical": "The final three churches were real congregations, but Adventist historicist interpretation also sees them as later phases of Christian history: post-Reformation formalism, the missionary and Advent awakening, and the end-time church in judgment-hour self-examination.",
        "literary": "The chapter completes Christ's searching diagnosis of the whole church. It moves from spiritual death, to faithful weakness, to lukewarm self-deception, and ends with Christ standing at the door.",
        "themes": ["Watchfulness", "Open door", "Laodicea", "Lukewarmness", "Book of life", "White raiment", "Eyesalve", "Christ's appeal"],
        "outline": [
            ("3:1-6", "Sardis", "A church with a living name is called to wake up, strengthen what remains, and walk in white."),
            ("3:7-13", "Philadelphia", "A weak but faithful church receives an open door and the promise of secure belonging with God."),
            ("3:14-22", "Laodicea", "A self-sufficient church is rebuked in love and invited to receive Christ's counsel and fellowship."),
        ],
        "periods": {
            "Sardis": "Adventist historicist application commonly connects Sardis with much of post-Reformation Protestantism, where inherited truth could harden into reputation without living vitality.",
            "Philadelphia": "Philadelphia is commonly connected with the era of revival, mission, and the Advent awakening, especially the open-door movement toward Christ's heavenly ministry.",
            "Laodicea": "Laodicea is commonly connected with the end-time church after 1844, when the judgment-hour message calls for repentance from self-sufficiency and renewed fellowship with Christ.",
        },
    },
    4: {
        "title": "The Heavenly Throne Room",
        "summary": "Revelation 4 lifts John from the messages to the churches into the heavenly throne room. Before seals are opened or judgments unfold, the church sees that God reigns, heaven worships, and creation exists for His pleasure.",
        "historical": "The churches of Asia lived under visible imperial pressure, but Revelation 4 relocates reality around heaven's throne. This is not escapism; it is the theological vantage point from which the rest of the book must be read.",
        "literary": "The chapter is a transition from the seven churches to the sealed scroll. It prepares Revelation 5 by presenting the throne, heavenly attendants, living creatures, elders, worship, and the Creator's worthiness.",
        "themes": ["Throne", "Worship", "Creation", "Seven lamps", "Living creatures", "Twenty-four elders", "Heavenly sanctuary"],
        "outline": [
            ("4:1", "The opened door", "John is invited through an opened heavenly door to see what must take place."),
            ("4:2-3", "The throne and the One seated", "The throne is the center of heaven's vision, surrounded by radiant glory and covenant mercy."),
            ("4:4-8", "Heavenly attendants", "Elders, lamps, sea, and living creatures surround the throne in royal-priestly worship."),
            ("4:9-11", "Creator worship", "The living creatures and elders worship God as Creator and sovereign Lord."),
        ],
        "periods": {
            "Throne room": "Adventist interpretation reads this as a heavenly sanctuary scene that prepares the reader for Christ's ministry, the scroll, and the opening of the seals.",
        },
    },
    5: {
        "title": "The Lamb and the Sealed Scroll",
        "summary": "Revelation 5 shows the sealed scroll in God's hand and the search for one worthy to open it. The Lion appears as the slain Lamb, takes the scroll, receives worship, and is praised as Redeemer of a worldwide kingdom of priests.",
        "historical": "The throne-room vision answers earthly uncertainty with heaven's verdict: history is held by God and is opened by the crucified and risen Christ.",
        "literary": "Chapter 5 completes the throne-room scene of chapter 4. Creation worship is joined by redemption worship, and the Lamb's worthiness prepares for the seals in chapter 6.",
        "themes": ["Lamb", "Scroll", "Worthy", "Redemption", "New song", "Prayers of the saints", "Kingdom of priests"],
        "outline": [
            ("5:1-5", "The sealed scroll", "No one can open the scroll until the Lion of Judah is announced as worthy."),
            ("5:6-8", "The slain Lamb", "John sees the Lion as a slain Lamb who takes the scroll and receives worship."),
            ("5:9-10", "The new song", "The Lamb is worthy because His blood redeemed a worldwide people and made them a kingdom of priests."),
            ("5:11-14", "Universal doxology", "Heaven and creation worship the Lamb with the One on the throne."),
        ],
        "periods": {
            "Lamb and scroll": "Adventist interpretation sees the Lamb's taking of the scroll as the authority scene that grounds the unfolding of history and later judgment, not as a denial of Christ's ongoing heavenly ministry.",
        },
    },
    6: {
        "title": "The First Six Seals",
        "summary": "Revelation 6 shows the Lamb opening the first six seals. The chapter moves from the four horsemen through the cry of the martyrs to cosmic signs and the great question: who shall be able to stand?",
        "historical": "The seals are read in Adventist historicism as a broad sequence through Christian history, parallel in some ways to the seven churches: apostolic advance, persecution, compromise, spiritual death, martyr witness, and signs pointing toward Christ's appearing.",
        "literary": "The chapter flows out of Revelation 5. The Lamb who is worthy now opens the seals, showing that even painful history is not outside His authority.",
        "themes": ["Seals", "Four horsemen", "Martyrs", "Cosmic signs", "Wrath of the Lamb", "Who can stand"],
        "outline": [
            ("6:1-8", "The first four seals", "The horsemen portray conquest, bloodshed, scarcity, and death as the Lamb opens history."),
            ("6:9-11", "The fifth seal", "The martyrs cry for vindication and are told to rest until the witness is complete."),
            ("6:12-17", "The sixth seal", "Cosmic signs and human terror bring the chapter to the question of who can stand."),
        ],
        "periods": {
            "First four seals": "Adventist historicist application commonly sees the white horse as apostolic purity and gospel conquest, the red horse as persecution, the black horse as compromise and spiritual scarcity, and the pale horse as deepened spiritual death.",
            "Fifth seal": "The fifth seal is commonly connected with the faithful martyrs whose blood cries for vindication during the long history of persecution.",
            "Sixth seal": "The sixth seal is commonly connected with signs associated with the approach of the Second Coming, while the final question points directly to Revelation 7.",
        },
    },
    7: {
        "title": "The Sealed Servants and the Great Multitude",
        "summary": "Revelation 7 answers the question at the end of chapter 6. Before final judgment is released, God's servants are sealed; John hears the 144,000 and sees a great multitude standing before the throne and the Lamb.",
        "historical": "The chapter is an interlude between the sixth and seventh seals. Adventist interpretation gives it special weight for sealing, final preparation, spiritual Israel, Sabbath-shaped allegiance, and the character of God's end-time people.",
        "literary": "The chapter moves from restraint, to sealing, to the numbered tribes, to the innumerable multitude, to the final assurance that the Lamb shepherds His people.",
        "themes": ["Sealing", "144,000", "Spiritual Israel", "Great multitude", "White robes", "Great tribulation", "Living fountains"],
        "outline": [
            ("7:1-4", "The sealing command", "The winds are restrained until God's servants are sealed in their foreheads."),
            ("7:5-8", "The tribes of the sealed", "The 144,000 are presented as ordered spiritual Israel, with the character of the group more important than speculation."),
            ("7:9-12", "The great multitude", "A countless international company worships before the throne and the Lamb."),
            ("7:13-17", "Those out of great tribulation", "The redeemed have washed their robes in the Lamb's blood and are shepherded forever by Him."),
        ],
        "periods": {
            "Sealing": "Adventist interpretation reads the sealing as God's final work of ownership, settled allegiance, and preservation before the winds of final crisis are released.",
            "144,000": "Adventist interpretation treats the 144,000 as the final sealed company from spiritual Israel, while acknowledging that Adventist writers may differ on whether the number itself is literal or symbolic.",
            "Great multitude": "The great multitude shows the final triumph of grace: God's redeemed people stand before the throne because salvation comes from God and from the Lamb.",
        },
    },
}


PERICOPE_RANGES = {
    2: [(1, 7, "Ephesus"), (8, 11, "Smyrna"), (12, 17, "Pergamos"), (18, 29, "Thyatira")],
    3: [(1, 6, "Sardis"), (7, 13, "Philadelphia"), (14, 22, "Laodicea")],
    4: [(1, 1, "The opened door"), (2, 3, "The throne and the One seated"), (4, 8, "Heavenly attendants"), (9, 11, "Creator worship")],
    5: [(1, 5, "The sealed scroll"), (6, 8, "The slain Lamb"), (9, 10, "The new song"), (11, 14, "Universal doxology")],
    6: [(1, 8, "First four seals"), (9, 11, "Fifth seal"), (12, 17, "Sixth seal")],
    7: [(1, 4, "Sealing"), (5, 8, "144,000"), (9, 12, "Great multitude"), (13, 17, "Those out of great tribulation")],
}


VERSE_INSIGHTS = {
    "2:1": "Christ begins with Ephesus by presenting Himself as the One who holds the seven stars and walks among the golden lampstands. The church's witness and its messengers are under His care and inspection, so the message is not administrative advice from a distant Lord but the searching word of the present Christ.",
    "2:2": "Christ commends Ephesus for works, labor, patience, and discernment. This was a church that could test false apostles and refuse evil. The verse honors doctrinal vigilance while preparing for the warning that truth defended without love becomes spiritually dangerous.",
    "2:3": "Ephesus had borne hardship for Christ's name and had not fainted. The church's endurance was real, not symbolic. Yet the verse sits beside the rebuke about lost love, showing that perseverance and affection must remain joined in Christian witness.",
    "2:4": "The rebuke is precise: Ephesus has left its first love. Christ does not deny its orthodoxy, labor, or endurance; He exposes the loss of warm devotion to Himself and generous love toward others.",
    "2:5": "Christ's remedy is remember, repent, and do the first works. The warning about removing the lampstand shows that a church's witness is not guaranteed by history or doctrine if living love is lost.",
    "2:6": "The deeds of the Nicolaitanes are rejected by Ephesus and by Christ. Their error is best understood as grace-abuse or compromise that separates profession from obedience, allowing practices that corrupt worship and discipleship.",
    "2:7": "The Spirit's appeal widens the message beyond Ephesus, and the promise of the tree of life reaches from Eden to the New Jerusalem. Christ calls the loveless but discerning church back toward restored life with God.",
    "2:8": "Smyrna hears from the First and the Last, the One who was dead and is alive. The title is fitted to a suffering church, because Christ speaks as the Lord who has already entered death and conquered it.",
    "2:9": "Christ knows Smyrna's tribulation and poverty, yet calls the church rich. The phrase synagogue of Satan names hostile opposition to Christ's people, but it must be handled carefully and never turned into contempt for Jewish people.",
    "2:10": "Smyrna is told not to fear suffering, prison, testing, or even death. The ten days mark a limited trial under God's boundary, and the crown of life promises victory beyond persecution.",
    "2:11": "The promise that the overcomer will not be hurt by the second death answers Smyrna's danger directly. The first death may come through persecution, but final death has no claim over those who belong to Christ.",
    "2:12": "Pergamos is addressed by the One with the sharp twoedged sword. This title fits a church that needs Christ's judging word to separate faithful confession from tolerated compromise.",
    "2:13": "Christ knows that Pergamos dwells where Satan's seat is, yet He also knows the faithful witness of Antipas. External pressure is real, but the example of Antipas proves that loyalty is possible even in hostile surroundings.",
    "2:14": "The doctrine of Balaam points back to Israel's seduction into idolatry and immorality. In Pergamos the danger is accommodation: teaching that makes compromise with the surrounding worship culture seem acceptable.",
    "2:15": "The doctrine of the Nicolaitanes reappears in Pergamos. What Ephesus hated as deeds is now tolerated as doctrine, showing how compromise hardens when a church gives theological cover to disobedience.",
    "2:16": "Christ commands the church to repent. The sword of His mouth will oppose the compromising teachers if the congregation refuses to correct what it has tolerated.",
    "2:17": "The hidden manna contrasts with compromised meals linked to idolatry, while the white stone and new name promise acceptance and identity from Christ. The overcomer receives deeper fellowship than the world can offer.",
    "2:18": "Thyatira hears from the Son of God, whose eyes are like fire and whose feet are like fine brass. The church faces internal seduction, so Christ appears as the One who sees and stands with unyielding holiness.",
    "2:19": "Christ commends Thyatira's charity, service, faith, patience, and growing works. The praise is real, but it does not cancel the rebuke that follows; active ministry cannot excuse tolerated corruption.",
    "2:20": "Jezebel represents prophetic pretension joined to idolatry and immorality. Whether an actual teacher or a symbolic name, the issue is influence that leads Christ's servants away from covenant faithfulness.",
    "2:21": "Christ gave Jezebel space to repent, which means even severe warnings are surrounded by mercy. Her refusal turns divine patience into a witness against settled rebellion.",
    "2:22": "The threatened judgment on Jezebel and those who commit adultery with her is severe because the danger is severe. The aim is still repentance before shared compromise becomes shared consequence.",
    "2:23": "Christ searches the reins and hearts and gives according to works. Works do not replace grace; they reveal allegiance, exposing whether faith is living, compromised, or rebellious.",
    "2:24": "Christ distinguishes the rest in Thyatira who have not accepted this doctrine or the so-called depths of Satan. False depth is unmasked wherever spiritual claims loosen loyalty to Christ.",
    "2:25": "The faithful are told to hold fast what they already have until Christ comes. The verse values perseverance in received truth rather than novelty for its own sake.",
    "2:26": "The overcomer keeps Christ's works unto the end and receives authority over the nations. True authority is promised to those who refuse corrupt influence now and remain faithful to Christ.",
    "2:27": "The rod of iron echoes messianic rule from Psalm 2. The promise is not cruel ambition but participation in Christ's final victory over powers that resist God's kingdom.",
    "2:28": "The morning star is a promise of Christ Himself and the dawn of His kingdom. Against Thyatira's darkness, the overcomer receives the true light.",
    "2:29": "The final appeal to hear the Spirit keeps the message from remaining ancient correspondence. Every church is summoned to learn from Christ's searching word to all the churches.",
    "3:1": "Sardis is addressed by Christ as the One who has the seven Spirits and the seven stars. A church with a name for life needs the fullness of the Spirit because Christ says its reputation hides spiritual death.",
    "3:2": "Christ commands Sardis to watch and strengthen what remains. The church is not told everything is lost, but what remains is ready to die unless awakened by repentance.",
    "3:3": "Sardis must remember, hold fast, and repent. The thief-language warns that a sleeping church will be surprised by Christ's visitation because it has confused past inheritance with present life.",
    "3:4": "Even in Sardis there are a few names who have not defiled their garments. Christ sees faithful individuals inside a declining community and promises they will walk with Him in white.",
    "3:5": "The overcomer receives white raiment, a secure name in the book of life, and Christ's confession before the Father and angels. The promise joins purity, assurance, and heavenly acknowledgment.",
    "3:6": "The repeated call to hear the Spirit makes Sardis a warning for every church with reputation but fading life. The message searches all who rely on a name rather than living faith.",
    "3:7": "Philadelphia hears from the Holy and True One who has the key of David. Christ's authority to open and shut is decisive, which comforts a church with little visible strength.",
    "3:8": "Christ sets an open door before Philadelphia because the church has little strength but has kept His word and not denied His name. Weakness does not prevent mission when Christ opens the way.",
    "3:9": "Philadelphia faces the synagogue of Satan, but Christ promises vindication. Opponents will know that Christ has loved His faithful people, so the church's future rests in His public acknowledgment.",
    "3:10": "Because Philadelphia kept the word of Christ's patience, He promises preserving care in the hour of temptation. The promise is about faithful endurance under Christ's keeping, not curiosity about escape.",
    "3:11": "Christ says He comes quickly and commands Philadelphia to hold fast. The crown is guarded by perseverance; what Christ gives must not be surrendered through neglect.",
    "3:12": "The overcomer becomes a pillar in the temple of God and receives the names of God, the New Jerusalem, and Christ. The promise answers weakness and instability with permanent belonging.",
    "3:13": "The Spirit's appeal closes the message to Philadelphia and makes its open door, little strength, and promised permanence instruction for all churches.",
    "3:14": "Laodicea is addressed by the Amen, the faithful and true Witness, and the beginning of God's creation. A self-deceived church needs Christ's reliable testimony more than its own assessment.",
    "3:15": "Christ knows Laodicea's works and finds the church neither cold nor hot. The problem is not honest weakness but useless complacency that neither refreshes nor heals.",
    "3:16": "Laodicea's lukewarmness makes Christ say He will spue it out of His mouth. The unpleasant image is severe mercy, meant to awaken a church comfortable with spiritual mediocrity.",
    "3:17": "Laodicea claims wealth and need of nothing, while Christ sees wretchedness, misery, poverty, blindness, and nakedness. The city's material symbols become a mirror of spiritual bankruptcy.",
    "3:18": "Christ counsels Laodicea to buy gold tried in the fire, white raiment, and eyesalve. Each gift answers a specific lack: genuine faith and love, Christ-given righteousness, and spiritual discernment.",
    "3:19": "Christ rebukes and chastens those He loves, so Laodicea's sharp message is not abandonment. The command is zeal and repentance because lukewarmness requires awakened desire.",
    "3:20": "Christ stands at the door and knocks. In context He is outside a self-sufficient church, offering restored fellowship to anyone who hears His voice and opens.",
    "3:21": "The overcomer is promised a seat with Christ on His throne, just as Christ overcame and sat with His Father. The highest promise is given to the church with the deepest need.",
    "3:22": "The seven messages end with the appeal to hear the Spirit. After lovelessness, persecution, compromise, corruption, dead reputation, faithful weakness, and lukewarmness, Christ still calls the churches to listen.",
    "4:1": "The opened door in heaven moves John from the churches to the throne room, and the trumpet-like voice summons him into prophetic vision. The command to come up is an invitation to see from heaven's perspective, not a symbol that removes the church from tribulation.",
    "4:2": "John is in the Spirit and immediately sees a throne set in heaven. The throne comes before every later conflict, teaching that divine sovereignty is the starting point for interpretation.",
    "4:3": "The jasper, sardius, and emerald rainbow communicate majesty, fiery brilliance, and covenant mercy. God is not described in ordinary bodily detail; His rule is shown through radiance and promise.",
    "4:4": "The twenty-four elders sit on thrones clothed in white raiment and crowned with gold. Their identity is debated, but the scene is royal, priestly, and worshipful, with all authority derivative from God.",
    "4:5": "Lightnings, thunderings, voices, and seven lamps of fire come from before the throne. The scene joins Sinai-like majesty with the seven Spirits, the fullness of the Spirit in a sanctuary setting.",
    "4:6": "The sea of glass and four beasts, better understood as living creatures, surround the throne. Their eyes signal fullness of perception in service of worship.",
    "4:7": "The lion, calf, human face, and flying eagle portray the fullness of created life before God. The living creatures represent creation in its strength, service, intelligence, and swiftness of praise.",
    "4:8": "The living creatures are full of eyes and never cease saying Holy, holy, holy. Heaven's worship centers on God's holiness, almighty power, and eternal being.",
    "4:9": "The living creatures give glory, honor, and thanks to the One on the throne. Worship is the truthful response of creation to the God who lives forever.",
    "4:10": "The elders fall down and cast their crowns before the throne. Heaven's highest servants do not cling to honor; they return all authority to God.",
    "4:11": "The chapter closes with the Creator's worthiness. God receives glory, honor, and power because He created all things and all things exist by His will.",
    "5:1": "John sees a scroll in the right hand of the One on the throne, written within and on the backside and sealed with seven seals. The future of history is held by God before it is opened by the Lamb.",
    "5:2": "A strong angel asks who is worthy to open the scroll and loose its seals. The question is moral and redemptive, not merely intellectual.",
    "5:3": "No one in heaven, earth, or under the earth can open or look upon the scroll. Creation cannot unlock its own destiny apart from the victory God provides in Christ.",
    "5:4": "John weeps much because no one is found worthy. His grief shows that the unopened scroll would mean history without disclosed resolution, judgment, or redemptive completion.",
    "5:5": "One elder tells John not to weep because the Lion of Judah, the Root of David, has prevailed. Messianic victory is announced before it is seen.",
    "5:6": "John turns and sees not a beast of force but a Lamb as though slain, standing in the midst of the throne. The seven horns and seven eyes, identified with the seven Spirits, show complete authority and Spirit-filled mission, while Christ's conquest is defined by sacrifice and resurrection.",
    "5:7": "The Lamb comes and takes the scroll from the right hand of the One on the throne. This act displays Christ's authority to unfold history because He has prevailed.",
    "5:8": "When the Lamb takes the scroll, the living creatures and elders fall before Him with harps and golden bowls of incense, which are the prayers of saints. Heaven treats the prayers of God's people as precious.",
    "5:9": "The new song declares the Lamb worthy because He was slain and redeemed people by His blood from every tribe, tongue, people, and nation.",
    "5:10": "The redeemed are made kings and priests to God and will reign on the earth. Salvation creates vocation, worship, and future authority under Christ.",
    "5:11": "John hears countless angels around the throne, the living creatures, and the elders. The worship of the Lamb expands from the inner circle to innumerable heavenly hosts.",
    "5:12": "The angels proclaim the slain Lamb worthy to receive power, riches, wisdom, strength, honor, glory, and blessing. Heaven's praise is total because Christ's redemption is total.",
    "5:13": "Every creature joins the doxology to the One on the throne and to the Lamb. The worship of God and the Lamb is shared without rivalry.",
    "5:14": "The living creatures say Amen and the elders fall down and worship. The chapter ends not with explanation but with adoration.",
    "6:1": "The Lamb opens the first seal and one living creature speaks with a thunder-like voice. The chapter begins in heaven, under the authority of the Lamb, before its effects unfold on earth.",
    "6:2": "The white horse, bow, crown, and conquering rider are read in Adventist historicism as the pure apostolic gospel advancing victoriously. The imagery is conquest, but its first application is gospel triumph rather than coercion.",
    "6:3": "The Lamb opens the second seal and the second living creature calls, Come and see. The heavenly command shows that the next movement of history is still under divine disclosure.",
    "6:4": "The red horse takes peace from the earth and brings a great sword. Adventist historicism connects this with persecution and bloodshed, especially the church under pagan Rome.",
    "6:5": "The black horse carries balances, shifting the imagery from bloodshed to scarcity and weighing. The seal points toward spiritual famine and compromise after the earlier purity and persecution.",
    "6:6": "The measured wheat and barley show scarcity, while the oil and wine are not hurt. The details suggest judgment restrained by mercy and a distinction between spiritual loss and what God preserves.",
    "6:7": "The fourth seal is opened and the fourth living creature calls. The sequence continues the deterioration from purity, to bloodshed, to scarcity, toward deathly spiritual conditions.",
    "6:8": "The pale horse is named Death, and hell follows. Adventist historicism commonly sees here the deep spiritual deadness associated with the medieval period, while the imagery also portrays the destructive results of sin.",
    "6:9": "The fifth seal reveals souls under the altar who were slain for the word of God and their testimony. The altar imagery presents martyrdom in sacrificial and sanctuary terms.",
    "6:10": "The martyrs cry, How long? Their plea is not revengeful hatred but a longing for God's truthful judgment and vindication.",
    "6:11": "White robes are given and the martyrs are told to rest a little season. God honors their witness while delaying final vindication until the testimony is complete.",
    "6:12": "The sixth seal opens with a great earthquake, a darkened sun, and a bloodlike moon. Adventist interpretation has linked these signs with historical portents pointing toward the nearness of Christ's coming.",
    "6:13": "The stars fall like figs shaken by a mighty wind. The cosmic language intensifies the sense that the created order itself announces approaching judgment.",
    "6:14": "The heaven departs as a scroll and every mountain and island is moved. The symbols portray the collapse of ordinary security before the presence of God.",
    "6:15": "All classes of humanity, from kings to slaves, hide themselves. Social rank gives no protection when the Lamb's judgment is revealed.",
    "6:16": "The terrified call for rocks and mountains to hide them from the face of the One on the throne and from the wrath of the Lamb. The gentlest title, Lamb, becomes fearful to those who rejected His mercy.",
    "6:17": "The chapter ends with the great question: who shall be able to stand? Revelation 7 answers by showing the sealed servants and the redeemed multitude.",
    "7:1": "Four angels hold the four winds so that harm does not yet fall on earth, sea, or trees. Final crisis is restrained by God until His servants are sealed.",
    "7:2": "Another angel ascends from the east with the seal of the living God. The eastward imagery suggests hope and divine initiative, and the seal marks ownership and loyalty.",
    "7:3": "The command is not to hurt the earth, sea, or trees until God's servants are sealed in their foreheads. The forehead points to settled conviction and character rather than outward appearance alone.",
    "7:4": "John hears the number of the sealed: 144,000 from the tribes of Israel. Adventist interpretation reads them as the final sealed company from spiritual Israel, with character more important than speculation over the number.",
    "7:5": "The list begins with Judah, then Reuben and Gad. Judah's first place points to the tribe of the Messiah and signals that the list is theological rather than an ordinary genealogy.",
    "7:6": "Aser, Nepthalim, and Manasses continue the ordered roster. The repeated number from each tribe emphasizes completeness and divine ordering.",
    "7:7": "Simeon, Levi, and Issachar are named in the same measured pattern. The list portrays God's people as known, counted, and arranged by Him.",
    "7:8": "Zabulon, Joseph, and Benjamin complete the tribal list. The omission of Dan and the unusual order show that the roster should be read symbolically and spiritually, not as a simple ethnic census.",
    "7:9": "John sees a great multitude no one can number from every nation, kindred, people, and tongue, clothed in white and holding palms. The vision expands from numbered order to global victory.",
    "7:10": "The multitude cries that salvation comes from God and the Lamb. No redeemed person claims credit for surviving tribulation or standing before the throne.",
    "7:11": "All the angels, elders, and living creatures fall before the throne and worship. The redeemed multitude's praise is joined by heaven's full worshiping order.",
    "7:12": "The sevenfold doxology gives blessing, glory, wisdom, thanksgiving, honor, power, and might to God forever. Salvation leads to worship that returns everything to God.",
    "7:13": "One elder asks who the white-robed multitude are and where they came from. The question prepares the interpretation of the vision so the reader does not guess.",
    "7:14": "The multitude came out of great tribulation and washed their robes in the Lamb's blood. Victory is not self-purification; it is cleansing through Christ joined to faithful endurance.",
    "7:15": "Because they are cleansed, they stand before God's throne and serve Him day and night in His temple. Their final security is worshipful nearness to God.",
    "7:16": "They hunger no more, thirst no more, and are no longer struck by sun or heat. The sufferings associated with wilderness, persecution, and deprivation are ended.",
    "7:17": "The Lamb in the midst of the throne shepherds them to living fountains of waters, and God wipes away every tear. The chapter's final answer to who can stand is pastoral: those shepherded by the Lamb.",
}


CROSS_REFERENCES = {
    2: ["Revelation 1:12-20", "Acts 20:28-31", "Numbers 25", "1 Kings 16-21", "Psalm 2", "Revelation 22:2"],
    3: ["Zechariah 4:1-6", "Isaiah 22:22", "Matthew 24:42-44", "Revelation 14:12", "Revelation 21:2", "Revelation 22:16"],
    4: ["Isaiah 6:1-4", "Ezekiel 1", "Daniel 7:9-10", "Exodus 19", "Zechariah 4:1-6", "Psalm 104"],
    5: ["Genesis 49:9-10", "Isaiah 53", "Daniel 7:13-14", "Psalm 141:2", "Exodus 19:5-6", "Revelation 6:1"],
    6: ["Zechariah 1:8-11", "Zechariah 6:1-8", "Matthew 24", "Joel 2:10-31", "Revelation 7:1-4", "Revelation 14:12"],
    7: ["Ezekiel 9:4-6", "Daniel 7:2", "Revelation 6:17", "Revelation 14:1-5", "Galatians 3:29", "Isaiah 49:10"],
}


SYMBOLS = {
    2: [
        ("Nicolaitanes", ["Revelation 2:6", "Revelation 2:15"], "Compromising teaching or practice that separated profession from holy obedience."),
        ("Balaam", ["Revelation 2:14"], "A pattern of idolatrous and immoral compromise that corrupts covenant loyalty."),
        ("Jezebel", ["Revelation 2:20"], "A symbol of religious influence that claims prophetic authority while leading God's servants into unfaithfulness."),
        ("Tree of life", ["Revelation 2:7"], "Restored access to the life with God lost in Eden and fulfilled in the New Jerusalem."),
    ],
    3: [
        ("White raiment", ["Revelation 3:5", "Revelation 3:18"], "Christ-given purity, victory, and righteousness."),
        ("Book of life", ["Revelation 3:5"], "The heavenly register of those who belong to Christ."),
        ("Key of David", ["Revelation 3:7"], "Christ's decisive royal authority to open and shut."),
        ("Eyesalve", ["Revelation 3:18"], "Spiritual discernment that heals Laodicean blindness."),
    ],
    4: [
        ("Throne", ["Revelation 4:2"], "God's sovereign rule over heaven, earth, worship, and history."),
        ("Twenty-four elders", ["Revelation 4:4"], "Heavenly royal-priestly worshipers whose authority is returned to God."),
        ("Seven lamps", ["Revelation 4:5"], "The fullness of the Spirit before the throne in sanctuary imagery."),
        ("Living creatures", ["Revelation 4:6-8"], "Creation's worshiping attendants around the throne."),
    ],
    5: [
        ("Sealed scroll", ["Revelation 5:1"], "God's held purpose for history, judgment, and redemption."),
        ("Lion of Judah", ["Revelation 5:5"], "The messianic conqueror promised from Judah and David's line."),
        ("Slain Lamb", ["Revelation 5:6"], "Christ conquering through sacrifice, resurrection, and redemptive worthiness."),
        ("Golden bowls", ["Revelation 5:8"], "The prayers of the saints treasured in heaven."),
    ],
    6: [
        ("White horse", ["Revelation 6:2"], "In Adventist historicism, the pure and conquering apostolic gospel."),
        ("Red horse", ["Revelation 6:4"], "Bloodshed and persecution in the church's historical experience."),
        ("Black horse", ["Revelation 6:5"], "Scarcity, compromise, and spiritual famine."),
        ("Pale horse", ["Revelation 6:8"], "Deathly spiritual condition and the destructive spread of sin."),
    ],
    7: [
        ("Four winds", ["Revelation 7:1"], "Restrained forces of final crisis and judgment."),
        ("Seal of God", ["Revelation 7:2-3"], "Divine ownership, settled allegiance, and preservation."),
        ("144,000", ["Revelation 7:4"], "The final sealed company from spiritual Israel."),
        ("Great multitude", ["Revelation 7:9"], "The redeemed from all nations standing before God and the Lamb."),
    ],
}


def pericope(chapter: int, verse: int) -> str:
    for start, end, name in PERICOPE_RANGES[chapter]:
        if start <= verse <= end:
            return name
    return CHAPTER_META[chapter]["title"]


def source_locator(chapter: int, verse: int) -> str:
    return f"{Path(DOCS[chapter]).name}: {pericope(chapter, verse)}"


def key_terms(text: str) -> list[str]:
    checks = [
        ("Nicolaitanes", "nicolait"),
        ("Balaam", "balaam"),
        ("Jezebel", "jezebel"),
        ("synagogue of Satan", "synagogue of satan"),
        ("hidden manna", "hidden manna"),
        ("white stone", "white stone"),
        ("morning star", "morning star"),
        ("seven Spirits", "seven spirits"),
        ("white raiment", "white raiment"),
        ("eyesalve", "eyesalve"),
        ("lukewarm", "lukewarm"),
        ("throne", "throne"),
        ("Lamb", "lamb"),
        ("seal", "seal"),
        ("beast", "beast"),
        ("144,000", "144"),
        ("great multitude", "great multitude"),
        ("temple", "temple"),
        ("tree of life", "tree of life"),
        ("book of life", "book of life"),
    ]
    lower = text.lower()
    return [label for label, needle in checks if needle in lower]


def technical_note(chapter: int, verse: int, text: str) -> str:
    lower = text.lower()
    if chapter == 5 and "seal" in lower:
        return "The sealed scroll represents history and destiny held in God's hand. It cannot be opened by human wisdom, political power, or angelic strength; it is entrusted to the Lamb because redemption gives Him the right to disclose and complete God's purpose."
    if "beast" in lower and chapter in {4, 5, 6, 7}:
        return "Here the KJV word beast is better understood as living creature. These beings are not hostile powers; they are worshiping attendants who surround the throne and give voice to creation's praise."
    if "nicolait" in lower:
        return "The Nicolaitanes represent a kind of compromise that separated Christian profession from holy obedience. The issue is not merely a strange ancient sect; it is the recurring temptation to make grace excuse practices that corrupt worship and discipleship."
    if "balaam" in lower:
        return "Balaam brings the Old Testament story of Numbers into Revelation. His name recalls a strategy of compromise in which idolatry and immorality weakened covenant loyalty from within."
    if "jezebel" in lower:
        return "Jezebel recalls the Old Testament queen who promoted Baal worship and persecuted God's faithful servants. Revelation uses the name for religious influence that claims authority while leading the church away from covenant faithfulness."
    if "synagogue of satan" in lower:
        return "The phrase synagogue of Satan names organized opposition to Christ's people, not contempt for Jewish people. Revelation is describing spiritual allegiance and hostility to the testimony of Jesus."
    if "hidden manna" in lower:
        return "The hidden manna contrasts with compromised meals and points to nourishment that comes from Christ Himself. The overcomer receives fellowship deeper than anything gained by accommodation to idolatry."
    if "white stone" in lower:
        return "The white stone and new name speak of acceptance, acquittal, and identity given by Christ. The believer's truest name is not assigned by hostile society but received from the Lord who knows His people."
    if "lukewarm" in lower:
        return "Lukewarmness is the condition of a church that has enough religion to feel secure but not enough surrender to be useful. Christ's severe language is meant to awaken love before self-deception becomes settled ruin."
    if "white raiment" in lower or "eyesalve" in lower or "gold tried" in lower:
        return "Gold, white raiment, and eyesalve answer Laodicea's false confidence. Christ offers tested faith, His own righteousness, and spiritual sight to a church that has mistaken prosperity and knowledge for living communion with Him."
    if "lamb" in lower:
        return "The Lamb language joins sacrifice, resurrection, authority, and worship. Revelation's victory is shaped by the cross before it is displayed in judgment."
    if "seal" in lower:
        return "Seal language points to ownership, settled allegiance, and preservation. In Revelation 6 the Lamb opens the seals of history; in Revelation 7 God seals His servants before final harm is released."
    if "throne" in lower:
        return "Throne language places every earthly crisis under God's government. Revelation teaches the church to interpret history from heaven's center, where worship, judgment, and mission are held together."
    if "144" in lower:
        return "The 144,000 are presented through Israel's tribal language, yet the unusual list and Revelation's symbolic style point toward spiritual Israel and final covenant loyalty rather than a simple ethnic census."
    if "great multitude" in lower:
        return "The great multitude shows the scope of redemption. The numbered, ordered people of God are also an innumerable company from every nation, gathered by the Lamb's blood and brought before God's throne."
    if chapter == 4:
        return "The background is the throne-vision language of Isaiah, Ezekiel, Daniel, and Sinai. The imagery teaches holiness, sovereignty, covenant mercy, and worship before Revelation moves into conflict."
    if chapter == 5:
        return "The Old Testament background includes Judah's Lion, David's Root, sacrificial lamb imagery, temple prayer, and the kingdom of priests. These strands explain why the Lamb is worthy to open what no creature can open."
    if chapter == 6:
        return "The seals use the language of prophetic horses, martyr witness, sanctuary imagery, and cosmic signs. The symbols are not detached puzzles; they are a theology of history under the authority of the Lamb."
    if chapter == 7:
        return "The chapter recalls Ezekiel's mark of protection, Daniel's winds of turmoil, and Israel's tribal language. Revelation gathers those themes around God's final people and the Lamb who brings them safely through."
    if chapter <= 3:
        return pericope_context_note(chapter, verse)
    return "The message should be heard within its immediate setting and within Revelation's larger concern for covenant loyalty, worship, witness, repentance, and hope."


def adventist_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter == 2:
        notes = {
            "Ephesus": "Across the wider movement of church history, Ephesus reflects the apostolic and immediately post-apostolic church: energetic in mission, careful with truth, and yet already in danger of losing the warmth of first love.",
            "Smyrna": "Across the wider movement of church history, Smyrna reflects the persecuted church under pagan Rome. Earthly poverty and martyrdom could not erase the riches Christ saw in His suffering people.",
            "Pergamos": "Across the wider movement of church history, Pergamos reflects the age when imperial favor made compromise attractive. The church's danger shifted from open persecution to religious accommodation with worldly power.",
            "Thyatira": "Across the wider movement of church history, Thyatira reflects the long medieval period, when corrupt religious influence grew while Christ still recognized works of love, faith, service, and perseverance among His faithful people.",
        }
        return notes[unit]
    if chapter == 3:
        notes = {
            "Sardis": "Across the wider movement of church history, Sardis reflects the post-Reformation danger of possessing a respected name while losing living spiritual power. Inherited truth could not substitute for watchfulness and repentance.",
            "Philadelphia": "Across the wider movement of church history, Philadelphia reflects the era of revival, mission, and Advent awakening. Its open door points beyond human strength to Christ's authority and His heavenly ministry.",
            "Laodicea": "Across the wider movement of church history, Laodicea reflects the judgment-hour church, especially after 1844. The issue is not lack of religious language but self-sufficiency in the very people called to bear Christ's final message.",
        }
        return notes[unit]
    if chapter == 4:
        return "The throne-room vision opens a sanctuary perspective. Before the seals unfold, heaven shows that God's government, Christ's ministry, and Creator worship stand behind the prophecy."
    if chapter == 5:
        return "The Lamb's taking of the scroll grounds the unfolding of history and judgment in Christ's victory. Prophecy is not controlled by earthly powers but by the crucified and risen Redeemer."
    if chapter == 6:
        notes = {
            "First four seals": "Read as a historical sequence, the first four seals trace a broad movement from apostolic purity and gospel advance, through persecution, into compromise, scarcity, and deepening spiritual death.",
            "Fifth seal": "Read as a historical sequence, the fifth seal gives voice to the martyrs whose witness calls for God's vindication. Their cry is held in heaven until the testimony of God's people is complete.",
            "Sixth seal": "Read as a historical sequence, the sixth seal includes signs associated with the nearness of Christ's coming while also moving toward the final day when every false security collapses before the Lamb.",
        }
        return notes[unit]
    notes = {
        "Sealing": "The sealing is God's final work of ownership, settled allegiance, and preservation before the winds of final crisis are released.",
        "144,000": "The 144,000 represent the final sealed company from spiritual Israel. The emphasis falls on character, allegiance, and God's ordering of His people more than on speculation.",
        "Great multitude": "The great multitude shows the final triumph of grace. Those who stand before the throne do so because salvation comes from God and from the Lamb.",
        "Those out of great tribulation": "This final scene answers the question of who can stand. The answer is a cleansed people, brought through tribulation and shepherded forever by the Lamb.",
    }
    return notes[unit]


def application_note(chapter: int, verse: int, text: str) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "The pastoral question is whether truth remains joined to love. Christ asks His people to preserve doctrine without becoming cold, to endure without becoming hard, and to receive correction as mercy rather than interruption.",
            "Smyrna": "The pastoral question is courage under pressure. Christ does not promise Smyrna an easy road, but He gives a promise large enough for suffering, loss, prison, and death.",
            "Pergamos": "The pastoral question is allegiance amid compromise. Christ calls His people to live where pressure is real without letting public advantage, fear, or desire weaken loyalty to Him.",
            "Thyatira": "The pastoral question is purity joined to perseverance. Christ values love, service, faith, and endurance, but He will not allow those graces to excuse teaching that corrupts worship.",
            "Sardis": "The pastoral question is whether reputation has replaced life. Christ calls His people to wakefulness, repentance, and renewed obedience before spiritual decline becomes settled.",
            "Philadelphia": "The pastoral question is faithful dependence. Christ opens doors for a church with little strength, teaching His people that mission rests on His authority rather than their visible power.",
            "Laodicea": "The pastoral question is surrender. Christ's severe diagnosis is spoken in love so that a self-satisfied church may receive His wealth, covering, sight, and fellowship.",
        }
        return notes[unit]
    if chapter == 4:
        return "The practical call is worship. Before believers try to interpret conflict, they must let the throne reorder fear, ambition, and imagination. Revelation begins heaven's central visions by teaching the church to look up."
    if chapter == 5:
        return "The practical call is trust and adoration. The future is opened by the Lamb who was slain, so prayer, mission, worship, and hope must remain centered on His redemptive victory rather than on human control."
    if chapter == 6:
        return "The practical call is patient faith. The seals do not make turmoil ultimate; they place turmoil under the Lamb's authority and teach believers to endure history without losing confidence in Christ."
    return "The practical call is preparation without panic. God knows His servants, marks them as His own, and brings them through tribulation into worship. The reader is invited to seek the Lamb's cleansing and settled allegiance now."


def pericope_context_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter == 2:
        notes = {
            "Ephesus": "Ephesus was a powerful commercial and religious city, known for the temple of Artemis and the pressure of public pagan life. In that setting, the lampstand image matters: the church exists to bear light in a difficult world, not merely to protect its own reputation.",
            "Smyrna": "Smyrna was prosperous and loyal to Rome, yet Christ's people there knew poverty, slander, and danger. The contrast between the city's wealth and the church's suffering helps explain why Christ measures riches by faithfulness rather than by social security.",
            "Pergamos": "Pergamos was marked by imperial loyalty and prominent pagan worship, so the pressure to blend Christian confession with civic religion was strong. Revelation treats that compromise as a matter of allegiance, not as harmless social convenience.",
            "Thyatira": "Thyatira was shaped by trade guilds and social networks where religious compromise could easily become part of ordinary economic life. Christ's warning shows that love and service are not allowed to excuse teaching that corrupts worship.",
        }
        return notes[unit]
    if chapter == 3:
        notes = {
            "Sardis": "Sardis had a reputation for past greatness and a history of falling because it failed to watch. That local background gives force to Christ's rebuke: a church can live on a name while its spiritual life is slipping away.",
            "Philadelphia": "Philadelphia had little strength, yet Christ speaks to it with royal authority. The open door is therefore not a reward for visible power but a promise that mission and access depend on Christ's key, not on human advantage.",
            "Laodicea": "Laodicea was known for wealth, textiles, and medical eye treatment, and Christ turns those local strengths into a spiritual diagnosis. The city could boast of resources, but the church needed the wealth, covering, and sight only Christ could give.",
        }
        return notes[unit]
    if chapter == 4:
        return "The throne-room scene draws on Isaiah, Ezekiel, Daniel, Sinai, and sanctuary imagery. Revelation shifts the reader from the pressure of earth to the worship of heaven so that later conflict is interpreted from God's rule rather than from fear."
    if chapter == 5:
        return "The chapter joins royal, sacrificial, and sanctuary imagery. The Lion promised from Judah appears as the slain Lamb; the prayers of the saints are treasured before God; and redemption creates a people who worship and serve as a kingdom of priests."
    if chapter == 6:
        return "The seals grow out of the Lamb's worthiness in chapter 5. The images of horses, altar, robes, and cosmic signs are not random disasters; they portray the church's history and the world's instability under the authority of Christ."
    return "Revelation 7 answers the question raised at the end of the sixth seal: who can stand? The answer is not human cleverness or fearlessness, but God's sealing work, the Lamb's blood, and a people brought safely into worship before the throne."


def pastoral_expansion_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "This keeps Ephesus from becoming merely a historical label. The same Christ who commends faithful labor also asks whether love still gives that labor its warmth, humility, and witness.",
            "Smyrna": "This keeps Smyrna from becoming only a story of ancient persecution. The suffering church still teaches believers to measure wealth by Christ's approval and to trust His life beyond the reach of death.",
            "Pergamos": "This keeps Pergamos from becoming only a lesson in ancient idolatry. It asks whether the church can recognize compromise when it arrives clothed in security, opportunity, or cultural respectability.",
            "Thyatira": "This keeps Thyatira from becoming only a medieval symbol. It asks whether love and service are being protected by truth, or whether tolerance has begun to excuse spiritual harm.",
            "Sardis": "This keeps Sardis from becoming only a warning about the past. It asks whether a church's name, history, and public reputation still correspond to living faith before Christ.",
            "Philadelphia": "This keeps Philadelphia from becoming merely a preferred church period. It asks whether weakness is being offered to Christ in faith, and whether the open door He gives is being entered with obedience.",
            "Laodicea": "This keeps Laodicea from becoming only a label for others. It asks whether Christ's loving rebuke is being received personally, before self-sufficiency shuts the door against fellowship.",
        }
        return notes[unit]
    if chapter == 4:
        return "This vision disciplines the imagination. It teaches believers that worship is not escape from reality but the truest way to see reality. The throne is not decorative scenery; it is the center from which the rest of Revelation must be understood."
    if chapter == 5:
        return "This keeps prophecy deeply Christ-centered. The scroll is not opened by calculation, fear, or force, but by the Lamb who was slain. The church can study what comes next only after learning to worship the One who is worthy."
    if chapter == 6:
        return "The seals should therefore produce sobriety without despair. They tell the truth about conquest, violence, scarcity, death, martyrdom, and final shaking, but they tell that truth from heaven, where the Lamb has already been declared worthy."
    return "The chapter does not invite speculation about status; it invites surrender to God's forming work. The sealed and the great multitude stand because they belong to God, have trusted the Lamb, and have been kept through the crisis by grace."


def christ_centered_expansion_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "The Savior among the lampstands cares not only that light burns, but that it burns with love. He commends what is real and calls His people back before witness becomes loveless duty.",
            "Smyrna": "The risen Christ speaks as One who has passed through death and lives forever. His presence gives suffering believers a hope stronger than the threats of empire or synagogue opposition.",
            "Pergamos": "Christ's sharp sword is not cruelty but mercy. His word cuts through mixed loyalties so that the church may be healed before compromise becomes its accepted way of life.",
            "Thyatira": "Christ searches the reins and hearts, so hidden influence is never hidden from Him. His eyes of fire expose corruption while His promise strengthens those who hold fast.",
            "Sardis": "Christ sees past reputation into reality. His rebuke is severe because He still calls the dying to wake, remember, repent, and walk with Him in white.",
            "Philadelphia": "Christ holds the key of David, so the faithful are not imprisoned by their weakness. He opens what no one can shut and secures those who keep His word.",
            "Laodicea": "The faithful Witness wounds pride in order to restore fellowship at the table. His knock is patient, personal, and full of mercy, showing that even severe rebuke is meant to bring the church back into communion.",
        }
        return notes[unit]
    if chapter == 4:
        return (
            "The vision also prepares the reader emotionally and spiritually. Judgment, conflict, and symbolic struggle will soon appear, but the first reality is not chaos; it is worship. "
            "The One on the throne is holy, eternal, and worthy, and His rule gives the church courage to face what follows."
        )
    if chapter == 5:
        return (
            "The scene also teaches how Revelation defines victory. Christ conquers as the slain Lamb, not by imitating the violence of earthly empires. "
            "His people are therefore called to overcome by faithfulness, witness, prayer, and worship, trusting that history is safest in the hands wounded for their redemption."
        )
    if chapter == 6:
        return (
            "That perspective keeps the symbols from becoming either cold chronology or frightening spectacle. The Lamb opens the seals, the martyrs are remembered, and the final question presses the reader toward surrender. "
            "The issue is not curiosity alone, but whether a person is being prepared to stand before Christ."
        )
    return (
        "That hope gives the chapter its pastoral warmth. The sealed are not merely counted; they are claimed. The multitude is not merely rescued; they are brought near. "
        "The final picture is a worshiping people whose security rests in God's name, Christ's blood, and the Shepherd who leads them to life."
    )


def final_completion_note(chapter: int, verse: int) -> str:
    if chapter <= 3:
        unit = pericope(chapter, verse)
        notes = {
            "Ephesus": "The appeal is not to choose between truth and love, but to let Christ restore both. A lampstand shines faithfully only when doctrine, endurance, and first love remain together.",
            "Smyrna": "The appeal is not to seek suffering, but to trust Christ when suffering comes. The crown of life belongs to those who cling to Him when earthly security is stripped away.",
            "Pergamos": "The appeal is to let Christ's word judge compromise before compromise reshapes the conscience. His correction preserves the church for faithful witness in a pressured world and before witness loses its clarity.",
            "Thyatira": "The appeal is to hold fast without making peace with corrupting influence. Christ treasures growing love and service, yet He calls that love to remain loyal and pure.",
            "Sardis": "The appeal is to wake before reputation becomes a substitute for life. Christ's warning is urgent because repentance can still strengthen what remains.",
            "Philadelphia": "The appeal is to keep Christ's word with patient confidence. Little strength is not failure when the door is opened by the One who holds the key.",
            "Laodicea": "The appeal is to open the door while Christ is still knocking. His rebuke is severe, but its goal is restored fellowship, healed sight, and a life no longer satisfied with itself.",
        }
        return notes[unit]
    if chapter == 4:
        return "A church that learns to worship here will read the rest of Revelation differently. It will see that every crisis is secondary to God's throne, every crown is borrowed, and every created thing finally exists for His pleasure."
    if chapter == 5:
        return "The reader is invited to let that worship shape interpretation. Revelation is not first about satisfying curiosity; it is about recognizing the Lamb's worthiness and joining heaven's confidence in His redemptive rule."
    if chapter == 6:
        return "The reader is therefore called to watchfulness rather than fear. The seals expose history's pain, but they also insist that Christ is not absent from history and will bring suffering witness to righteous completion."
    return "The reader is invited to receive that hope personally. God's final people are not secure because they understand every detail perfectly, but because they are marked by His ownership, washed by the Lamb, and drawn into worship."


def manuscript_note(chapter: int, verse: int, text: str) -> str:
    excerpt = select_doc_excerpt(chapter, verse, text)
    if not excerpt:
        return technical_note(chapter, verse, text)
    banned = [
        "This reading",
        "Adventist interpretation",
        "Other Adventist",
        "Commentators",
        "commentators",
        "non-Adventist",
        "A fair synthesis",
        "Key symbols",
    ]
    if any(item in excerpt for item in banned):
        return technical_note(chapter, verse, text)
    if words(excerpt) > 130:
        sentences = split_sentences(excerpt)
        excerpt = " ".join(sentences[:4])
    return excerpt or technical_note(chapter, verse, text)


def theological_note(chapter: int, verse: int) -> str:
    if chapter <= 3:
        unit = pericope(chapter, verse)
        notes = {
            "Ephesus": "Christ binds truth, love, and witness together. He praises faithful discernment, but He will not let orthodoxy excuse a heart that has cooled toward Him.",
            "Smyrna": "Christ binds suffering and hope together. He sees affliction from within resurrection life, so the church's weakness is held inside His victory over death.",
            "Pergamos": "Christ binds confession and obedience together. Loyalty to His name must not be separated from purity in worship, teaching, and daily allegiance.",
            "Thyatira": "Christ binds love and holiness together. Works may increase, but growth is not healthy if the church tolerates teaching that leads people away from covenant faithfulness.",
            "Sardis": "Christ binds reputation and reality together. A living name cannot replace a living walk, and unfinished works must be brought back before God.",
            "Philadelphia": "Christ binds weakness and mission together. The church's security rests in His key, His promise, and His knowledge of those who keep His word.",
            "Laodicea": "Christ binds rebuke and fellowship together. The Amen exposes self-sufficiency because He desires a church rich in faith, clothed in His righteousness, and open to His presence.",
        }
        return notes[unit]
    if chapter == 4:
        return "The throne steadies the chapter. Before Revelation shows conflict on earth, it shows worship in heaven, teaching that history must be read from God's sovereignty, holiness, and Creator-rights."
    if chapter == 5:
        return "The Lamb's worthiness governs the chapter. Revelation's prophecy is shaped by Christ's sacrifice, resurrection, authority, and redemption of a people who worship and serve God."
    if chapter == 6:
        return "The seals place history under the Lamb. They are not loose disasters; they are opened by Christ, who holds the church's witness, suffering, and final vindication within His authority."
    return "God's preserving ownership gives the chapter its strength. The sealed servants and the white-robed multitude show that final preparation, endurance, and salvation are grounded in the Lamb's cleansing work."


def background_bridge_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "That setting makes endurance costly and love necessary. The congregation had to bear light where public religion, trade, and civic pride constantly pressed believers toward easier loyalties.",
            "Smyrna": "That setting makes Christ's comfort especially tender. The church's outward vulnerability was not evidence of divine neglect, because heaven measured its life by faithfulness rather than status.",
            "Pergamos": "That setting explains why compromise could seem reasonable. The pressure was not only private temptation, but a public world where worship, loyalty, food, status, and safety were intertwined.",
            "Thyatira": "That setting shows how false teaching can become ordinary. When work, guild life, relationships, and worship overlap, compromise may arrive not as rebellion but as practical accommodation.",
            "Sardis": "That setting gives Christ's rebuke a sharper edge. A city known for reputation and careless confidence becomes the setting for a church warned that a name for life is not life itself.",
            "Philadelphia": "That setting makes the open door gracious rather than triumphant. The church has little strength, yet Christ's authority gives mission, security, and belonging that social power cannot supply.",
            "Laodicea": "That setting makes the diagnosis painfully concrete. Wealth, clothing, and eye medicine become spiritual metaphors because the church's greatest danger is mistaking resources for communion with Christ.",
        }
        return notes[unit]
    if chapter == 4:
        return "This background keeps the imagery from floating free. John is not inventing religious decoration; he is gathering biblical throne, temple, creation, and covenant language into one worshiping scene."
    if chapter == 5:
        return "This background explains why heaven's sorrow turns to worship. The problem is not lack of information, but lack of worthiness; only redemptive victory can open the scroll."
    if chapter == 6:
        return "This background keeps the seals sober rather than sensational. The symbols describe history under Christ's authority, where witness, suffering, judgment, and final accountability are all held together."
    return "This background keeps the sealing scene pastoral rather than speculative. Revelation is showing whom God claims, prepares, and brings safely into worship before the throne."


def historicist_bridge_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "That historical application does not cancel the local message. It shows how even a church near the freshness of apostolic witness can need renewal when zeal for truth begins to lose the warmth of love.",
            "Smyrna": "That historical application does not make suffering abstract. It lets Smyrna stand for every season when Christ's people appear poor and exposed, yet are rich because their life is hidden with Him.",
            "Pergamos": "That historical application explains why favor can be spiritually dangerous. A church may survive persecution and still be wounded when comfort, influence, and social acceptance reshape its obedience.",
            "Thyatira": "That historical application gives the warning long reach. It shows how religious corruption can grow over time while Christ still knows, preserves, and strengthens those who refuse its claims.",
            "Sardis": "That historical application makes the warning searching. Recovered truth must remain living truth, or a church may keep the name of reform while losing the watchful spirit that reform requires.",
            "Philadelphia": "That historical application makes the promise missionary. Christ's open door points to a season of revival, Scripture, mission, and Advent hope that rests on His authority rather than human strength.",
            "Laodicea": "That historical application makes the rebuke urgent for the judgment-hour church. The people called to give Christ's final message must not confuse doctrinal possession with living dependence on Christ.",
        }
        return notes[unit]
    if chapter == 4:
        return "This sanctuary perspective matters because the rest of Revelation will move through conflict, judgment, and final restoration. Heaven's throne gives those later scenes their center of gravity."
    if chapter == 5:
        return "This makes prophecy deeply Christ-centered. The future is not opened by calculation, fear, or force, but by the Lamb whose sacrifice gives Him the right to finish God's purpose."
    if chapter == 6:
        return "This reading should never become cold chronology. The seals expose the spiritual condition of history and press the reader toward the final question of who is prepared to stand."
    return "This reading should not turn the sealed into a subject for curiosity alone. The emphasis rests on settled allegiance, Christ's cleansing, and God's ability to keep His servants."


def sentence_candidates(texts: list[str]) -> list[str]:
    candidates: list[str] = []
    for text in texts:
        cleaned = public_guard(text)
        if not cleaned:
            continue
        pieces = [part.strip() for part in re.split(r"(?<=[.!?])\s+", cleaned) if len(part.strip().split()) >= 5]
        if not pieces and words(cleaned) >= 5:
            pieces = [cleaned]
        for sentence in pieces:
            sentence = public_guard(sentence)
            lowered = sentence.lower()
            if sentence and lowered not in [item.lower() for item in candidates]:
                candidates.append(sentence)
    return candidates


def opening_bridge_note(chapter: int, verse: int) -> str:
    unit = pericope(chapter, verse)
    if chapter <= 3:
        notes = {
            "Ephesus": "His presence keeps doctrinal courage and first love together, so hatred of false deeds does not become loveless religion.",
            "Smyrna": "His victory over death gives the suffering church a future larger than the threats placed before it.",
            "Pergamos": "His judging word is also a saving word, cutting through compromise so loyalty can be restored.",
            "Thyatira": "His searching gaze is mercy for the faithful and warning for every influence that corrupts worship.",
            "Sardis": "His rebuke is the voice of One who can still awaken what is dying and strengthen what remains.",
            "Philadelphia": "His authority turns weakness into witness because the open door belongs to Him.",
            "Laodicea": "His rebuke is severe because fellowship is still possible and the door may still be opened.",
        }
        return notes[unit]
    return theological_note(chapter, verse)


def fit_paragraph(base: str, additions: list[str], minimum: int, maximum: int) -> str:
    paragraph_sentences = sentence_candidates([base])
    if not paragraph_sentences:
        paragraph_sentences = sentence_candidates(additions[:1])
    paragraph: list[str] = []
    for sentence in paragraph_sentences:
        trial = " ".join(paragraph + [sentence])
        if words(trial) <= maximum:
            paragraph.append(sentence)
    for sentence in sentence_candidates(additions):
        if words(" ".join(paragraph)) >= minimum:
            break
        if sentence.lower() in " ".join(paragraph).lower():
            continue
        trial = " ".join(paragraph + [sentence])
        if words(trial) <= maximum:
            paragraph.append(sentence)
    if words(" ".join(paragraph)) > maximum:
        trimmed: list[str] = []
        for sentence in paragraph:
            trial = " ".join(trimmed + [sentence])
            if words(trial) <= maximum:
                trimmed.append(sentence)
        paragraph = trimmed
    return public_guard(" ".join(paragraph))


def compose_commentary(chapter: int, verse: int, text: str) -> str:
    insight = VERSE_INSIGHTS[f"{chapter}:{verse}"]
    context = manuscript_note(chapter, verse, text)
    prophetic = adventist_note(chapter, verse)
    application = application_note(chapter, verse, text)
    theology = theological_note(chapter, verse)
    technical = technical_note(chapter, verse, text)
    opening_additions = [christ_centered_expansion_note(chapter, verse), opening_bridge_note(chapter, verse)]
    prophetic_additions = [historicist_bridge_note(chapter, verse)]
    application_additions = [pastoral_expansion_note(chapter, verse)]
    if chapter >= 4:
        opening_additions.append(theology)
        prophetic_additions.extend([theology, background_bridge_note(chapter, verse)])
        application_additions.append(final_completion_note(chapter, verse))
    paragraphs = [
        fit_paragraph(
            insight,
            opening_additions,
            70,
            108,
        ),
        fit_paragraph(
            context,
            [technical, pericope_context_note(chapter, verse), background_bridge_note(chapter, verse)],
            70,
            112,
        ),
        fit_paragraph(
            prophetic,
            prophetic_additions,
            74,
            112,
        ),
        fit_paragraph(
            application,
            application_additions,
            74,
            110,
        ),
    ]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    result = "\n\n".join(paragraphs)
    if words(result) < 300:
        extras = [
            [opening_bridge_note(chapter, verse), christ_centered_expansion_note(chapter, verse), theology],
            [background_bridge_note(chapter, verse), technical],
            [historicist_bridge_note(chapter, verse), pericope_context_note(chapter, verse), theology],
            [pastoral_expansion_note(chapter, verse), final_completion_note(chapter, verse)],
        ]
        for index in sorted(range(len(paragraphs)), key=lambda item: words(paragraphs[item])):
            for sentence in sentence_candidates(extras[index]):
                if words(result) >= 300:
                    break
                if sentence.lower() in result.lower():
                    continue
                maximum = 110 if index == 3 else 112
                trial_paragraph = f"{paragraphs[index]} {sentence}".strip()
                if words(trial_paragraph) <= maximum:
                    paragraphs[index] = trial_paragraph
                    result = "\n\n".join(paragraphs)
            if words(result) >= 300:
                break
    if words(result) > 500:
        kept = []
        for paragraph in result.split("\n\n"):
            trial = "\n\n".join(kept + [paragraph])
            if words(trial) <= 500:
                kept.append(paragraph)
        result = "\n\n".join(kept)
    return public_guard(result)


def commentary_for(chapter: int, verse: int, text: str) -> tuple[str, dict]:
    insight = public_guard(VERSE_INSIGHTS[f"{chapter}:{verse}"])
    context = public_guard(manuscript_note(chapter, verse, text))
    technical = public_guard(technical_note(chapter, verse, text))
    adventist = public_guard(adventist_note(chapter, verse))
    application = public_guard(application_note(chapter, verse, text))
    result = compose_commentary(chapter, verse, text)

    layers = {
        "detailedExplanation": result,
        "exegesis": insight,
        "historicalBackground": context,
        "technicalNotes": technical,
        "adventistPropheticInsight": adventist,
        "propheticTimeline": adventist,
        "otherCommentaryInsights": public_guard(theological_note(chapter, verse)),
        "application": application,
        "reviewFlags": [],
    }
    return result, layers


def chapter_for_kjv(chapter: int) -> list[dict]:
    kjv = json.loads(KJV_PATH.read_text())
    found = next(item for item in kjv["chapters"] if int(item["chapter"]) == chapter)
    return found["verses"]


def build_audit(chapter: int, verse: int) -> dict:
    locator = source_locator(chapter, verse)
    doc = ref(DOC_SOURCE_IDS[chapter], locator, "manuscript-synthesis", 1)
    mcnulty = ref(MCNULTY_ID, "Priority Adventist commentary for Revelation", "adventist-interpretation", 1)
    shorter = ref(SHORTER_ID, "Technical and Old Testament background support", "technical-background", 5)
    return {
        "exegesis": [doc, mcnulty],
        "historicalBackground": [doc],
        "technicalNotes": [doc, shorter],
        "adventistPropheticInsight": [doc, mcnulty],
        "propheticTimeline": [doc, mcnulty],
        "otherCommentaryInsights": [shorter],
        "application": [doc, mcnulty],
    }


def build_chapter(chapter: int) -> dict:
    meta = CHAPTER_META[chapter]
    verses = []
    for entry in chapter_for_kjv(chapter):
        verse_num = int(entry["verse"])
        verse_ref = f"Revelation {chapter}:{verse_num}"
        text = entry["text"]
        detailed, layers = commentary_for(chapter, verse_num, text)
        audit = build_audit(chapter, verse_num)
        flat_sources = []
        seen = set()
        for source_list in audit.values():
            for item in source_list:
                if item["sourceId"] not in seen:
                    flat_sources.append(item)
                    seen.add(item["sourceId"])
        verses.append(
            {
                "verse": verse_ref,
                "bibleText": text,
                "explanation": layers["exegesis"],
                "historicalBackground": layers["historicalBackground"],
                "symbolicMeaning": layers["technicalNotes"],
                "adventistInsight": layers["adventistPropheticInsight"],
                "propheticSignificance": layers["propheticTimeline"],
                "danielConnection": "Daniel supplies key background for Revelation's symbolic world, especially kingdom, throne, beast, judgment, sealing, and faithful-endurance themes.",
                "crossReferences": CROSS_REFERENCES[chapter],
                "application": layers["application"],
                "sources": flat_sources,
                "commentary": layers,
                "sourceAudit": audit,
                "reviewStatus": "verified-seed",
            }
        )

    doc_ref = ref(DOC_SOURCE_IDS[chapter], Path(DOCS[chapter]).name, "chapter-manuscript", 1)
    outline = [{"range": r, "title": t, "summary": s} for r, t, s in meta["outline"]]
    return {
        "chapterNumber": chapter,
        "title": meta["title"],
        "summary": meta["summary"],
        "historicalContext": meta["historical"],
        "literaryContext": meta["literary"],
        "themes": meta["themes"],
        "outline": outline,
        "verses": verses,
        "symbols": [
            {
                "symbol": symbol,
                "references": references,
                "meaning": meaning,
                "sources": [doc_ref],
            }
            for symbol, references, meaning in SYMBOLS[chapter]
        ],
        "charts": [{"id": f"revelation-{chapter}-overview", "title": meta["title"], "type": "chapter-map"}],
        "images": [],
        "crossReferences": CROSS_REFERENCES[chapter],
        "danielConnections": [
            {
                "danielText": "Daniel's apocalyptic and judgment background",
                "revelationText": f"Revelation {chapter}",
                "sources": [doc_ref],
            }
        ],
        "teachingNotes": {
            "openingQuestion": "What does this chapter reveal about Jesus and His people before it reveals events?",
            "mainPoint": meta["summary"],
            "keyVerses": [f"Revelation {chapter}:{PERICOPE_RANGES[chapter][0][0]}", f"Revelation {chapter}:{PERICOPE_RANGES[chapter][-1][1]}"],
            "importantSymbols": [item[0] for item in SYMBOLS[chapter]],
            "discussionQuestions": [
                "What does Christ commend, warn, or reveal in this chapter?",
                "How does the chapter connect worship, faithfulness, and prophetic hope?",
                "What response should this chapter produce in a church class or personal study?",
            ],
            "commonMisunderstandings": [
                "Reading the symbols apart from the immediate passage and Old Testament background.",
                "Using prophetic interpretation without receiving the pastoral appeal.",
            ],
            "adventistEmphasis": next(iter(meta["periods"].values())),
            "closingAppeal": "Hear what the Spirit is saying and respond to Christ with repentance, worship, endurance, and hope.",
        },
        "evangelisticNotes": {
            "mainDoctrinalTheme": "Christ-centered prophecy forms a faithful people for worship, witness, and final hope.",
            "keyBibleTexts": [f"Revelation {chapter}:{PERICOPE_RANGES[chapter][0][0]}", f"Revelation {chapter}:{PERICOPE_RANGES[chapter][-1][1]}"],
            "flow": [
                "Begin with the immediate scene and what it reveals about Christ.",
                "Explain the main symbol or historical application carefully.",
                "Connect the passage to worship, faithfulness, and hope.",
                "Invite a practical response rather than fear or speculation.",
            ],
            "simpleIllustrations": ["A lamp only matters if it gives light; prophecy only helps if it leads people to trust and obey Christ."],
            "appealQuestion": "Will you receive Christ's message in this chapter as light for your own life?",
            "cautions": ["Avoid sensational language and unsupported details."],
            "sources": [doc_ref],
        },
        "reflectionQuestions": [
            "What picture of Jesus stands at the center of this chapter?",
            "Where does this chapter call me to repent, endure, worship, or hope?",
            "How should this chapter shape the way I teach Revelation to others?",
        ],
        "sources": [doc_ref, ref(MCNULTY_ID, "Priority Adventist commentary for Revelation", "adventist-interpretation", 1)],
    }


def update_resources() -> None:
    path = CONTENT / "resources" / "bibliography.json"
    data = json.loads(path.read_text())
    resources = {item["id"]: item for item in data.get("resources", [])}
    for chapter, doc_path in DOCS.items():
        resources[DOC_SOURCE_IDS[chapter]] = {
            "id": DOC_SOURCE_IDS[chapter],
            "title": doc_path.stem,
            "author": "User-provided manuscript",
            "type": "Word manuscript",
            "tradition": "Adventist",
            "interpretiveCategory": "Adventist historicist",
            "howUsed": "Internal manuscript source for synthesized chapter commentary and hidden audit metadata.",
            "citationFormat": "Internal manuscript metadata retained for audit only.",
        }
    resources.setdefault(
        MCNULTY_ID,
        {
            "id": MCNULTY_ID,
            "title": "Revelation: Practical Living in the Judgment Hour",
            "author": "Norman McNulty",
            "type": "Commentary",
            "tradition": "Adventist",
            "interpretiveCategory": "Adventist historicist",
            "howUsed": "Priority Adventist source for Revelation interpretation, sanctuary emphasis, and prophetic framework.",
            "citationFormat": "Internal source metadata retained for audit only.",
        },
    )
    resources.setdefault(
        SHORTER_ID,
        {
            "id": SHORTER_ID,
            "title": "Revelation: A Shorter Commentary",
            "author": "G. K. Beale and David Campbell",
            "type": "Commentary",
            "tradition": "Non-Adventist",
            "interpretiveCategory": "Academic / evangelical",
            "howUsed": "Technical, literary, and Old Testament background support where it does not control prophetic interpretation.",
            "citationFormat": "Internal source metadata retained for audit only.",
        },
    )
    data["resources"] = sorted(resources.values(), key=lambda item: item["id"])
    path.write_text(json.dumps(data, indent=2) + "\n")


def main() -> None:
    if not KJV_PATH.exists():
        raise FileNotFoundError(f"Missing KJV source: {KJV_PATH}")
    for chapter, path in DOCS.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing manuscript for Revelation {chapter}: {path}")
        headings = read_doc_headings(path)
        if len(headings) < 3:
            raise ValueError(f"Manuscript for Revelation {chapter} does not look like the expected study document.")
    update_resources()
    for chapter in range(2, 8):
        target = CONTENT / "revelation" / f"chapter-{chapter:02}.json"
        target.write_text(json.dumps(build_chapter(chapter), indent=2) + "\n")
    print("Imported Revelation manuscript commentary for chapters 2-7.")


if __name__ == "__main__":
    main()
